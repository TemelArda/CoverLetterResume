import docx
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


doc = docx.Document('Template.docx')
newDoc = docx.Document()

date = input('Enter Date: ')
posName = input('Enter PosName: ')
compName = input('Enter CompName: ')
compNameShort = input('Enter CompNameShort: ')
compLocation = input('Enter compLocation: ')



newPara = ""
fullText = []
for para in doc.paragraphs:
    cur = para.text
    for word in cur.split():
        if word == "DATE":
            newPara += date
        elif word == "PosName":
            newPara += posName
        elif word == "CompName":
            newPara += compName
        elif word == "CompNameShort":
            newPara += compNameShort
        elif word == "CompLocation":
            newPara += compLocation   
        else:
            newPara += word + ' '
    print(newPara)
    paragraph = newDoc.add_paragraph(newPara)
    paragraphFormat = paragraph.paragraph_format
    paragraphFormat.space_before = Pt(3)
    paragraphFormat.space_after = Pt(3)
    newPara = ""
newDoc.styles['Normal'].font.name = "Century Gothic"
newDoc.styles['Title'].font.size = Pt(24)
newDoc.styles['Title'].font.bold = True
newDoc.styles['Title'].font.color.rgb = RGBColor(0x0, 0x0, 0x0)
newDoc.paragraphs[0].style = "Title"

for section in newDoc.sections:
    section.left_margin = Inches(.8)
    section.right_margin = Inches(.8)
    section.top_margin = Inches(.7)
    section.bottom_margin = Inches(.7)

print(newDoc.styles['Normal'].font.name)


newDoc.save("CoverLetter_ArdaTemel"+compNameShort+".docx")
    


